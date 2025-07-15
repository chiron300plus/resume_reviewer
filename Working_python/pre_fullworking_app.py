import gradio as gr
from openai import OpenAI
import PyPDF2
import docx
import os
import json
from docx import Document
from datetime import datetime

# ✅ Set login credentials
USERNAME = "admin"
PASSWORD = "letmein123"  # Change this to your preferred password

# ✅ Initialize OpenAI
client = OpenAI(api_key="MY_API_KEY")  # Replace with your real key

# Extract text from uploaded resume
def extract_text(file_path):
    ext = os.path.splitext(file_path)[-1].lower()
    if ext == ".pdf":
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            return "\n".join([page.extract_text() for page in reader.pages])
    elif ext == ".docx":
        doc = docx.Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    return "❌ Unsupported file type."

# Feedback, rewritten resume, and interview questions
def get_resume_feedback_and_rewrite(file_path, job_description):
    try:
        resume_text = extract_text(file_path)
        prompt = f"""
You are a professional resume coach. Review the resume below and return your response strictly in this JSON format:
{{
  "scores": {{"grammar": 0-10, "structure": 0-10, "job_fit": 0-10}},
  "suggestions": ["..."],
  "rewritten_summary": "...",
  "improved_bullet_point": "...",
  "missing_keywords": ["..."],
  "rewritten_resume": "Full rewritten resume text here, with formatting and sections clearly labeled.",
  "mock_interview_questions": ["..."]
}}
Resume: {resume_text}
Target Job or Description: {job_description if job_description else "N/A"}
"""
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7
        )
        content = response.choices[0].message.content
        data = json.loads(content)

        scores = (
            f"📝 Grammar: {data['scores']['grammar']}/10\n"
            f"📐 Structure: {data['scores']['structure']}/10\n"
            f"🎯 Job Fit: {data['scores']['job_fit']}/10"
        )
        suggestions = "\n- " + "\n- ".join(data["suggestions"])
        summary = data["rewritten_summary"]
        bullet = data["improved_bullet_point"]
        keywords = ", ".join(data["missing_keywords"])
        rewritten_resume_text = data["rewritten_resume"]
        interview_questions = "\n- " + "\n- ".join(data["mock_interview_questions"])
        question_list = data["mock_interview_questions"]

        doc = Document()
        for line in rewritten_resume_text.splitlines():
            doc.add_paragraph(line)
        filename = f"Rewritten_Resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(filename)

        return scores, suggestions, summary, bullet, keywords, filename, resume_text, rewritten_resume_text, interview_questions, question_list
    except Exception as e:
        return [f"❌ Error: {str(e)}"] * 5 + [None, "", "", "", []]

# GPT feedback on audio answer
def get_audio_feedback(audio_path, selected_question):
    try:
        if not audio_path or not os.path.exists(audio_path):
            return "❌ Please record an answer first."
        with open(audio_path, "rb") as f:
            transcript = client.audio.transcriptions.create(model="whisper-1", file=f).text
        prompt = f"""
You are an interview coach. Here's a mock interview question and the candidate's spoken answer.

Question: "{selected_question}"
Answer: "{transcript}"

Give constructive feedback on content, clarity, and relevance. Be professional and concise.
"""
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}]
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"❌ Error: {e}"

# ✅ Login function
def check_login(username, password):
    if username == USERNAME and password == PASSWORD:
        return gr.update(visible=False), gr.update(visible=True), ""
    else:
        return gr.update(), gr.update(visible=False), "❌ Incorrect username or password."

# ✅ Full App UI with Login Protection
with gr.Blocks(title="🔐 Secure Resume App") as app:
    # Login section
    with gr.Column(visible=True) as login_section:
        gr.Markdown("### 🔐 Login to Access the Resume Tool")
        username_input = gr.Textbox(label="Username")
        password_input = gr.Textbox(label="Password", type="password")
        login_btn = gr.Button("🔓 Login")
        login_error = gr.Textbox(label="", interactive=False)

    # Main app section (hidden until login)
    with gr.Column(visible=False) as main_app:
        gr.Markdown("## 📄 Upload your resume and optionally paste a job description below.")
        with gr.Row():
            resume_file = gr.File(label="📄 Upload Resume (.pdf, .docx, .txt)", type="filepath")
            job_input = gr.Textbox(label="📋 Job Title or Full Job Description", lines=4)
        submit = gr.Button("🧠 Analyze Resume")
        with gr.Row():
            scores_out = gr.Textbox(label="📊 Scores")
            suggestions_out = gr.Textbox(label="🛠 Suggestions")
        with gr.Row():
            summary_out = gr.Textbox(label="✍️ Rewritten Summary")
            bullet_out = gr.Textbox(label="🔁 Improved Bullet Point")
            keywords_out = gr.Textbox(label="❗ Missing Keywords")
        file_out = gr.File(label="⬇️ Download Rewritten Resume (.docx)")
        gr.Markdown("## 🆚 Resume Comparison")
        with gr.Row():
            original_out = gr.Textbox(label="📄 Original Resume", lines=20)
            rewritten_out = gr.Textbox(label="✍️ Rewritten Resume", lines=20)
        gr.Markdown("## 🎤 Mock Interview Questions")
        interview_out = gr.Textbox(label="Interview Questions", lines=10)
        gr.Markdown("## 🎙️ Practice Interview Responses")
        with gr.Row():
            question_dropdown = gr.Dropdown(label="🧠 Select a Mock Question", choices=[], interactive=True)
            audio_input = gr.Audio(sources="microphone", type="filepath", label="🎤 Record Your Answer")
        feedback_btn = gr.Button("🔍 Get Feedback on Answer")
        feedback_out = gr.Textbox(label="🧠 GPT Feedback")

    # Link logic
    login_btn.click(check_login, [username_input, password_input], [login_section, main_app, login_error])
    submit.click(get_resume_feedback_and_rewrite, [resume_file, job_input], [
        scores_out, suggestions_out, summary_out, bullet_out,
        keywords_out, file_out, original_out, rewritten_out,
        interview_out, question_dropdown
    ])
    feedback_btn.click(get_audio_feedback, [audio_input, question_dropdown], feedback_out)

if __name__ == "__main__":
    print("✅ Launching app...")
    app.queue().launch(
        server_name="0.0.0.0",
        server_port=int(os.environ.get("PORT", 8080))
    )
